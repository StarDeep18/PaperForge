"""
DOCX Parser using python-docx.

Extracts text and metadata from Microsoft Word documents.
"""

from pathlib import Path
from typing import Optional

import docx

from app.core.logging import logger
from app.domain.entities.document import DocumentMetadata
from app.domain.exceptions import DocumentProcessingError


class DOCXParseResult:
    """Result of parsing a DOCX document."""

    def __init__(self, text: str, metadata: DocumentMetadata):
        self.text = text
        self.metadata = metadata
        self.page_breaks: dict[int, int] = {}  # DOCX doesn't have reliable page breaks


class DOCXParser:
    """
    Extracts text and metadata from DOCX files using python-docx.

    Handles:
    - Full text extraction from paragraphs and tables
    - Core properties metadata extraction
    - Heading detection for section tracking
    """

    def parse(self, file_path: str) -> DOCXParseResult:
        """
        Parse a DOCX file and extract text + metadata.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            DOCXParseResult with text and metadata.
        """
        path = Path(file_path)
        if not path.exists():
            raise DocumentProcessingError("", f"File not found: {file_path}")

        try:
            doc = docx.Document(str(path))
        except Exception as e:
            raise DocumentProcessingError("", f"Failed to open DOCX: {e}") from e

        # Extract text from paragraphs
        text_parts: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        # Extract metadata from core properties
        props = doc.core_properties
        metadata = DocumentMetadata(
            title=props.title if props.title else None,
            authors=[props.author] if props.author else [],
            keywords=props.keywords.split(",") if props.keywords else [],
            word_count=len(full_text.split()),
            page_count=0,  # DOCX doesn't reliably store page count
        )

        logger.info(
            f"Parsed DOCX: {path.name}, "
            f"paragraphs={len(doc.paragraphs)}, "
            f"words={metadata.word_count}"
        )

        return DOCXParseResult(text=full_text, metadata=metadata)
